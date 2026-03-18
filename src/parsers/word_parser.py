"""
Word文档解析器
支持.doc和.docx格式
"""

from pathlib import Path
from typing import List, Optional

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import olefile
    OLE_AVAILABLE = True
except ImportError:
    OLE_AVAILABLE = False

from .base import BaseParser
from ..core.types import ParsedDocument, Table
from ..core.exceptions import DocumentParseError


class WordParser(BaseParser):
    """Word文档解析器"""

    SUPPORTED_EXTENSIONS = [".docx", ".doc"]

    def _do_parse(self, file_path: str) -> ParsedDocument:
        """解析Word文件"""
        path = Path(file_path)

        if path.suffix.lower() == ".docx":
            return self._parse_docx(file_path)
        else:
            return self._parse_doc(file_path)

    def _parse_docx(self, file_path: str) -> ParsedDocument:
        """解析.docx格式"""
        if not DOCX_AVAILABLE:
            raise DocumentParseError("python-docx未安装，请运行: pip install python-docx")

        doc = docx.Document(file_path)
        metadata = self._extract_metadata(Path(file_path))

        # 提取文本
        text_content = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_content.append(para.text)

        # 提取表格
        tables = []
        for table in doc.tables:
            parsed_table = self._convert_docx_table(table)
            if parsed_table:
                tables.append(parsed_table)

        full_text = "\n".join(text_content)

        return ParsedDocument(
            text=full_text,
            tables=tables,
            metadata=metadata,
            file_path=file_path,
            file_type="docx"
        )

    def _parse_doc(self, file_path: str) -> ParsedDocument:
        """
        解析.doc格式（旧版Word）
        注：.doc格式解析较复杂，这里提供基础实现
        """
        # 尝试使用antiword或textract等工具
        # 这里提供一个简化的实现
        text_content = []

        try:
            # 尝试读取为文本（很多.doc文件实际上包含可提取的文本）
            with open(file_path, 'rb') as f:
                content = f.read()
                # 尝试提取文本（简单启发式方法）
                text = self._extract_text_from_doc(content)
                if text:
                    text_content.append(text)
        except Exception as e:
            raise DocumentParseError(
                f"解析.doc文件失败: {file_path}",
                details={"error": str(e), "suggestion": "建议将.doc转换为.docx格式"}
            )

        metadata = self._extract_metadata(Path(file_path))

        return ParsedDocument(
            text="\n".join(text_content),
            tables=[],
            metadata=metadata,
            file_path=file_path,
            file_type="doc"
        )

    def _extract_text_from_doc(self, content: bytes) -> str:
        """从.doc文件中提取文本（启发式方法）"""
        # 尝试解码为文本
        try:
            # 查找文本段落（Word文档中通常以特定标记分隔）
            text_parts = []
            i = 0
            while i < len(content):
                # 查找可打印字符序列
                if 32 <= content[i] <= 126 or content[i] in [0x0a, 0x0d]:
                    start = i
                    while i < len(content) and (32 <= content[i] <= 126 or content[i] in [0x0a, 0x0d, 0x20]):
                        i += 1
                    chunk = content[start:i]
                    try:
                        text = chunk.decode('utf-8', errors='ignore')
                        if len(text) > 3:  # 过滤短片段
                            text_parts.append(text)
                    except:
                        pass
                else:
                    i += 1

            # 合并并清理文本
            text = " ".join(text_parts)
            # 清理多余空白
            text = " ".join(text.split())
            return text
        except:
            return ""

    def _convert_docx_table(self, table) -> Optional[Table]:
        """将docx表格转换为Table对象"""
        if not table.rows:
            return None

        # 提取表头
        headers = []
        if table.rows:
            headers = [cell.text.strip() for cell in table.rows[0].cells]

        # 提取数据行
        rows = []
        for row in table.rows[1:]:
            row_data = [cell.text.strip() for cell in row.cells]
            if any(row_data):
                rows.append(row_data)

        return Table(
            headers=headers,
            rows=rows
        )
